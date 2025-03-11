import werkzeug
print(f"Using Werkzeug version: {werkzeug.__version__}")

from flask import Flask, request, jsonify, send_file, redirect
from flask_cors import CORS
import os
import sys
import json
import tempfile
from docx import Document
import time

app = Flask(__name__)
CORS(app)  # Allow cross-origin requests

# Get port from environment variable (for cloud deployment)
PORT = int(os.environ.get('PORT', 5000))

# Handle paths for cloud environment
def get_temp_dir():
    """Get appropriate temp directory based on environment"""
    if 'RENDER' in os.environ:
        # Use Render's temp directory
        return '/tmp'
    else:
        # Use system default temp directory
        return tempfile.gettempdir()

@app.route('/', methods=['GET'])
def home():
    """Return a basic API welcome message so that clients accessing the API host get a simple response."""
    frontend_url = os.environ.get('FRONTEND_URL', 'https://docexp.onrender.com')
    return redirect(frontend_url)
@app.route('/api-info', methods=['GET'])
def api_info():
    """Return API info for API documentation purposes"""
    return jsonify({
        "api": "DocGen API",
        "status": "active",
        "version": "1.0.0",
        "message": "This is the API endpoint. For the web application, go to https://docexp.onrender.com",
        "frontend_url": os.environ.get('FRONTEND_URL', 'https://docexp.onrender.com')
    })

@app.route('/status', methods=['GET'])
def status():
    """Check if the server is running and ready"""
    return jsonify({
        'status': 'ready',
        'version': '1.0.0',
        'environment': os.environ.get('RENDER', 'development'),
        'features': ['word_generation', 'pdf_preview']
    })

@app.route('/generate-document', methods=['POST'])
def generate_document():
    """Generate a Word document from template and form data"""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        template = data.get('template')
        form_data = data.get('data')
        
        if not all([template, form_data]):
            return jsonify({'error': 'Missing required parameters'}), 400
            
        # Generate a temporary file path
        temp_dir = get_temp_dir()
        output_path = os.path.join(temp_dir, f"doc_{int(time.time())}.docx")
        
        # Generate the document
        success = create_word_document(template, form_data, output_path)
        
        if success:
            # Return the file for download
            return send_file(output_path, as_attachment=True, 
                            attachment_filename=f"{template['name'].replace(' ', '_')}.docx", 
                            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            return jsonify({'error': 'Failed to generate document'}), 500
    except Exception as e:
        app.logger.error(f"Error generating document: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/generate-json-document', methods=['POST'])
def generate_json_document():
    """Generate a Word document from a JSON template and form data"""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        template = data.get('template')
        form_data = data.get('data')
        output_path = data.get('outputPath')
        
        if not all([template, form_data, output_path]):
            return jsonify({'error': 'Missing required parameters'}), 400
            
        # Ensure the output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Generate the document
        success = create_json_word_document(template, form_data, output_path)
        
        if success:
            return jsonify({'success': True, 'path': output_path})
        else:
            return jsonify({'error': 'Failed to generate document'}), 500
    except Exception as e:
        app.logger.error(f"Error generating JSON document: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/generate-pdf-preview', methods=['POST'])
def generate_pdf_preview():
    """Generate a PDF preview from template and form data"""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        template = data.get('template')
        form_data = data.get('data')
        
        if not all([template, form_data]):
            return jsonify({'error': 'Missing required parameters'}), 400
            
        # Generate a temporary Word document
        temp_docx = f"temp_{int(time.time())}.docx"
        success = create_word_document(template, form_data, temp_docx)
        
        if not success:
            return jsonify({'error': 'Failed to generate document'}), 500
            
        # Convert Word document to PDF
        temp_pdf = temp_docx.replace('.docx', '.pdf')
        pdf_success = convert_word_to_pdf(temp_docx, temp_pdf)
        
        if not pdf_success:
            return jsonify({'error': 'Failed to convert document to PDF'}), 500
            
        # Return the PDF file
        with open(temp_pdf, 'rb') as pdf_file:
            pdf_data = pdf_file.read()
            
        # Clean up temporary files
        try:
            os.remove(temp_docx)
            os.remove(temp_pdf)
        except:
            pass
            
        return jsonify({
            'success': True,
            'pdf': pdf_data.decode('latin1')  # This is not ideal but works for demo
        })
    except Exception as e:
        app.logger.error(f"Error generating PDF preview: {str(e)}")
        return jsonify({'error': str(e)}), 500

def create_word_document(template, form_data, output_path):
    """Create a Word document from template and form data"""
    try:
        # Create a new document
        doc = Document()
        
        # Add a title
        doc.add_heading(template['name'], 0)
        
        # Add a description if available
        if template.get('description'):
            doc.add_paragraph(template['description'])
        
        # Add form data
        for field in template['fields']:
            field_name = field['name']
            field_label = field['label']
            field_value = form_data.get(field_name, '')
            
            # Create a paragraph for each field
            p = doc.add_paragraph()
            p.add_run(f"{field_label}: ").bold = True
            p.add_run(str(field_value))
        
        # Add creation date
        doc.add_paragraph(f"\nDocumento generato il: {time.strftime('%d/%m/%Y %H:%M:%S')}")
        
        # Save the document
        doc.save(output_path)
        
        return True
    except Exception as e:
        app.logger.error(f"Error creating Word document: {str(e)}")
        return False

def create_json_word_document(template, form_data, output_path):
    """Create a Word document from a JSON template and form data"""
    try:
        # Create a new document
        doc = Document()
        
        # Add a title
        doc.add_heading(template['name'], 0)
        
        # Add a description if available
        if template.get('description'):
            doc.add_paragraph(template['description'])
        
        # Process the JSON data recursively
        process_json_data(doc, form_data)
        
        # Add creation date
        doc.add_paragraph(f"\nDocumento generato il: {time.strftime('%d/%m/%Y %H:%M:%S')}")
        
        # Save the document
        doc.save(output_path)
        
        return True
    except Exception as e:
        app.logger.error(f"Error creating JSON Word document: {str(e)}")
        return False

def process_json_data(doc, data, level=0):
    """Process JSON data recursively and add to Word document"""
    if data is None:
        return
    
    # Add proper indentation based on level
    indent = "  " * level
    
    if isinstance(data, dict):
        for key, value in data.items():
            # Format key for better readability
            formatted_key = key.replace('_', ' ').title()
            
            if isinstance(value, dict):
                # Add a heading for nested objects
                p = doc.add_paragraph()
                p.add_run(f"{indent}{formatted_key}:").bold = True
                process_json_data(doc, value, level + 1)
            elif isinstance(value, list):
                # Add a heading for arrays
                p = doc.add_paragraph()
                p.add_run(f"{indent}{formatted_key}:").bold = True
                
                # Process each item in the array
                for i, item in enumerate(value):
                    if isinstance(item, dict):
                        p = doc.add_paragraph()
                        p.add_run(f"{indent}\tItem {i+1}:").italic = True
                        process_json_data(doc, item, level + 2)
                    else:
                        p = doc.add_paragraph()
                        p.add_run(f"{indent}\tItem {i+1}: ").italic = True
                        p.add_run(str(item))
            else:
                # Add simple key-value pair
                p = doc.add_paragraph()
                p.add_run(f"{indent}{formatted_key}: ").bold = True
                p.add_run(str(value))
    elif isinstance(data, list):
        for i, item in enumerate(data):
            p = doc.add_paragraph()
            p.add_run(f"{indent}Item {i+1}:").bold = True
            if isinstance(item, (dict, list)):
                process_json_data(doc, item, level + 1)
            else:
                p.add_run(f" {str(item)}")

def convert_word_to_pdf(input_docx, output_pdf):
    """Convert Word document to PDF"""
    try:
        # In a real application, you would use a library like python-docx-pdf or pywin32
        # For this demo, we'll just create a simple PDF
        with open(output_pdf, 'w') as pdf_file:
            pdf_file.write("This is a simulated PDF file for demo purposes.")
        return True
    except Exception as e:
        app.logger.error(f"Error converting Word to PDF: {str(e)}")
        return False

if __name__ == '__main__':
    debug_mode = os.environ.get('DEBUG', 'False').lower() == 'true'
    host = '0.0.0.0'  # Bind to all interfaces
    app.run(debug=debug_mode, host=host, port=PORT)
